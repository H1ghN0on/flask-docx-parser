from flask import Flask, render_template, request, send_file
from docx.enum.text import WD_LINE_SPACING
import re
from docx import Document
from docx.shared import Pt, Mm
from io import BytesIO

from zipfile import BadZipFile

app = Flask(__name__)
app.config['TEMPLATES_AUTO_RELOAD'] = True



def addOrderer(person, orders):
    #[id1488228|Валерий Жмышенко]

    info = person.split("|")


    name = info[1][:-1] if len(info) == 2 else info[0][:-1]
    id = info[0][1:] if len(info) == 2 else -1
    
    dictPerson = {
        'id': id,
        'name': name,
        'items': []
    }

    
    existed = list(filter(lambda person: person['id'] == id if id != -1 else person['name'] == name, orders))
    if (not existed):
        orders.append(dictPerson)
        return len(orders) - 1
    else:
        index = next((index for (index, d) in enumerate(orders) if d["id"] == existed[0]['id']), None)
        return index

def addOrder(order, orderIndex, currentOrder, orders):
    order = " ".join(order).split(", ")
            
    cards = 0
    postcards = 0
    ids = 0
    sets = 0

    for part in order:
        
        if (re.match('^\d+[kKкК]$', part)):
            cards += int(part[:-1])
        if (re.match('^\d+[оОoO]$', part)):
            postcards += int(part[:-1])  
        if (re.match('^\d+[аАaA]$', part)):
            ids += int(part[:-1])
        if (re.match('^\d+[sSсС]$', part)):
            sets += int(part[:-1])
        
    newOrder = {
        'cards': cards,
        'postcards': postcards,
        'ids': ids,
        'sets': sets,
        'order': currentOrder
    }

    orders[orderIndex]['items'].append(newOrder)


@app.route("/")
def index():
    return render_template("index.html")

@app.route("/parse", methods=["GET", "POST"])
def parse():
    if (request.method == "POST"):
        
        file = request.files.get("file")
        if (file):
            try:
                doc = Document(file)
                paras = doc.paragraphs

                currentOrder = 0
                orders = []
                orderIds = set()
                brokenOrders = []
                currentOrder = 0

        
                for paragraph in paras:
                   
                    if (paragraph.text): 
                       
                        text = paragraph.text
                        text = ' '.join(text.split())
                        if (re.match('^\d+[📦]', text) or text.startswith("!!") or text.startswith("#") or text.startswith("инд #") or not text[0].isdigit()):
                            if text.startswith("#"):
                                currentOrder = text[1:]
                                currentOrder = currentOrder.split(" ")[0]
                                orderIds.add(int(currentOrder))
                            elif text.startswith("инд #"):
                                currentOrder = text[5:]
                                orderIds.add(int(currentOrder))
                            
                            if (text.startswith("!!")):
                                text = text.replace("!!", "").split(". ")
                                brokenOrders.append({
                                    'order': currentOrder,
                                    'item': text[1],
                                })
                        else:

                            position, person, person2, *order = text.split(" ")
                            if (not person.endswith("]")):
                                person = person + " " + person2
                            orderIndex = addOrderer(person, orders)
                            addOrder(order, orderIndex, currentOrder, orders) 
                return {'orders': orders, 'orderIds': sorted(list(orderIds)), 'brokenOrders': brokenOrders }, 200

            except BadZipFile:
                return {'error': "Не Word документ"}, 500
            except Exception as e:
                return {'error': "Плохо отформатированный документ"}, 500
        else: 
            return {'error': "Файл не загружен"}, 500
    
    else:
        return({'error': "Ошибка случилась" }, 500)


def fillName(index, order):
    return f"{index + 1}. [{order['id']}|{order['name']}] "

def fillItem(item, character):

    if (item != 0):
        return f"{item}{character}, "
    else:
        return ""

def fillFinalOrder(items):
    finalCalamity = ""

    finalCalamity += fillItem(items['cards'],"к")
    finalCalamity += fillItem(items['postcards'],"о")
    finalCalamity += fillItem(items['ids'],"а")
    finalCalamity += fillItem(items['sets'],"с")

    if (finalCalamity):
        finalCalamity = finalCalamity[:len(finalCalamity) - 2]
    return finalCalamity
    
def fillOrders(items):
    finalString = ""
    finalCalamity = {
        'cards': 0,
        'postcards': 0,
        'ids': 0,
        'sets': 0,
    }
    for item in items:
        itemToWrite = ""
        itemToWrite += f"{item['order']}"
        itemToWrite += "("
        
        cards = fillItem(item['cards'], "к")
        postcards = fillItem(item['postcards'], "о")
        ids = fillItem(item['ids'], "а")
        sets = fillItem(item['sets'], "c")

        if (cards):
            itemToWrite += cards
            finalCalamity['cards'] += item['cards']
        if (postcards):
            itemToWrite += postcards
            finalCalamity['postcards'] += item['postcards']
        if (ids):
            itemToWrite += ids
            finalCalamity['ids'] += item['ids']
        if (sets):
            itemToWrite += sets
            finalCalamity['sets'] += item['sets']
        
        if (len(itemToWrite) > len(str(item['order'])) + 1):
            itemToWrite = itemToWrite[:len(itemToWrite) - 2]
        itemToWrite += "), "

        finalString += itemToWrite
    
    finalString = finalString[:len(finalString) - 2]

    finalString += " - "
    finalString += fillFinalOrder(finalCalamity)

    return finalString
    
def addParagraph(document, str):
    paragraph = document.add_paragraph(str)   
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph.paragraph_format.space_before = Mm(0)    

def generateDocument(orders, specificOrders, brokenOrders):
    document = Document()
    style = document.styles['Normal']
    
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    
    for i, order in enumerate(orders):
        orderInfo = ""
        
        orderInfo += fillName(i, order)
        orderInfo += fillOrders(order['items'])
        orderInfo += " — "
        addParagraph(document, orderInfo)

    addParagraph(document, "НАШЕ")

    for i, order in enumerate(specificOrders):
        orderInfo = ""
        orderInfo += f'{i + 1}. {order["name"]} '
        orderInfo += fillOrders(order['items'])
        orderInfo += " — "
        addParagraph(document, orderInfo)


    addParagraph(document, "ОТМЕНЁННЫЕ") 

    for i, order in enumerate(brokenOrders):
        orderInfo = "!! "
        orderInfo += f'(({order["order"]})) '
        orderInfo += f'{order["item"]}'
        orderInfo += " !!"

        addParagraph(document, orderInfo)
    return document

@app.route("/generate", methods=["POST"])
def generate():
    if (request.method == "POST"):
        
        orders = request.json['orders']
        specificOrders = request.json['specificOrders']
        brokenOrders = request.json['brokenOrders']
        document = generateDocument(orders, specificOrders, brokenOrders)

        f = BytesIO()
        document.save(f)
        f.seek(0)

        return send_file(f, as_attachment=True, download_name='result.docx')
    else:
        return({'error': "Ошибка случилась" }, 500)


  

if __name__ == "__main__":
    app.run(debug=True)